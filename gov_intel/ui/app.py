"""The main GOV.UK Policy Intelligence Workstation window."""

from __future__ import annotations

import json
import logging
import os
import re
import threading
import webbrowser
from datetime import datetime
from pathlib import Path
from typing import Callable

import requests
import tkinter as tk
from tkinter import colorchooser, filedialog, messagebox, scrolledtext, simpledialog, ttk

from .. import ai_engine, archive, config
from ..app_state import PROFILES_FILE, AppState
from ..gov_api import classify_attachment_url, deep_harvest_gov_uk, sanitize_filename
from ..models import Document, normalize_keyword_rules
from ..storage import save_json
from .data_grid import DataGridViewerWindow
from .pdf_viewer import PDFViewerWidget
from .split_compare import SplitPDFComparatorWindow

try:
    import docx
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False

logger = logging.getLogger(__name__)

PALETTE_PRESETS = [
    ("Yellow", [1.0, 0.9, 0.2], "#ffef33"),
    ("Red", [1.0, 0.35, 0.35], "#ff5959"),
    ("Green", [0.3, 0.85, 0.4], "#4dd866"),
    ("Blue", [0.3, 0.65, 1.0], "#4da6ff"),
    ("Purple", [0.75, 0.4, 0.95], "#bf66f2"),
    ("Orange", [1.0, 0.6, 0.2], "#ff9933"),
    ("Teal", [0.2, 0.8, 0.8], "#33cccc"),
    ("Pink", [1.0, 0.5, 0.75], "#ff80bf"),
]


def _attachment_label(url: str) -> str:
    tail = url.rstrip("/").rsplit("/", 1)[-1].replace("%20", " ")
    return f"{classify_attachment_url(url)}: {tail or url}"


def _rgb_to_hex(rgb: list[float]) -> str:
    """Converts 0.0-1.0 float RGB values to a Hex color string."""
    if not rgb or len(rgb) < 3:
        return "#ffef33"
    r = int(max(0.0, min(1.0, rgb[0])) * 255)
    g = int(max(0.0, min(1.0, rgb[1])) * 255)
    b = int(max(0.0, min(1.0, rgb[2])) * 255)
    return f"#{r:02x}{g:02x}{b:02x}"


# Continuous heatmap gradient stops (fraction-of-max -> RGB 0-255), used to
# color-code keyword-density cells in the Policy Analytics matrix so relative
# intensity is visible at a glance instead of just three flat buckets.
_HEATMAP_STOPS: list[tuple[float, tuple[int, int, int]]] = [
    (0.0, (255, 251, 235)),   # faint amber - just above zero
    (0.20, (254, 240, 138)),  # pale yellow
    (0.45, (250, 204, 21)),   # yellow
    (0.70, (249, 115, 22)),   # orange
    (1.0, (185, 28, 28)),     # deep red - hottest
]


def _heat_color(t: float) -> tuple[str, str]:
    """Maps a 0.0-1.0 intensity fraction to a (background_hex, text_hex) pair
    along a white->yellow->orange->red gradient, choosing readable text color
    based on the background's perceived luminance."""
    t = max(0.0, min(1.0, t))
    for (t0, c0), (t1, c1) in zip(_HEATMAP_STOPS, _HEATMAP_STOPS[1:]):
        if t0 <= t <= t1:
            local_t = 0.0 if t1 == t0 else (t - t0) / (t1 - t0)
            r = round(c0[0] + (c1[0] - c0[0]) * local_t)
            g = round(c0[1] + (c1[1] - c0[1]) * local_t)
            b = round(c0[2] + (c1[2] - c0[2]) * local_t)
            luminance = (0.299 * r + 0.587 * g + 0.114 * b) / 255
            fg = "#0f172a" if luminance > 0.55 else "#ffffff"
            return f"#{r:02x}{g:02x}{b:02x}", fg
    r, g, b = _HEATMAP_STOPS[-1][1]
    return f"#{r:02x}{g:02x}{b:02x}", "#ffffff"


class GovApp:
    def __init__(self, root: tk.Tk):
        self.root = root
        self.root.title("🇬🇧 GOV.UK Policy Intelligence Workstation")
        self.root.geometry("1400x950")

        config.ensure_data_dirs()
        self.state = AppState.load()

        self.active_docs: list[Document] = []
        self.selected_doc: Document | None = None
        self.active_atts: list[str] = []

        self.fav_selected_doc_id: str | None = None
        self.fav_active_atts: list[str] = []
        self.all_fav_attachments: dict[str, list[str]] = {"pdf": [], "link": [], "data": []}

        self.style = ttk.Style()
        self.style.theme_use("clam")

        self.nb = ttk.Notebook(root)
        self.nb.pack(fill="both", expand=True, padx=8, pady=8)
        self.tab_control = ttk.Frame(self.nb)
        self.tab_reader = ttk.Frame(self.nb)
        self.tab_favs = ttk.Frame(self.nb)
        self.tab_kw = ttk.Frame(self.nb)
        self.tab_analytics = ttk.Frame(self.nb)

        self.nb.add(self.tab_control, text="🛰️ Control Panel")
        self.nb.add(self.tab_reader, text="📋 Intelligence Reader")
        self.nb.add(self.tab_favs, text="⭐ Favorites Hub")
        self.nb.add(self.tab_kw, text="🧠 Keyword Brain")
        self.nb.add(self.tab_analytics, text="📈 Policy Analytics")

        self._build_control_tab()
        self._build_reader_tab()
        self._build_favs_tab()
        self._build_kw_tab()
        self._build_analytics_tab()
        self.refresh_fav_hub()

        # The Keyword Brain should be "straight in there" with the full picture:
        # activate the aggregated Master Overview profile as the active working
        # set every time the app starts, rather than requiring a manual click.
        self.load_master_overview_profile(silent=True)

    # ======================================================================
    # Control tab
    # ======================================================================
    def _build_control_tab(self) -> None:
        f = ttk.Frame(self.tab_control, padding=15)
        f.pack(fill="both", expand=True)
        f.columnconfigure(1, weight=1)
        f.rowconfigure(10, weight=1)

        ttk.Label(f, text="GOV.UK Advanced Intelligence Platform", font=("Segoe UI", 16, "bold"),
                  foreground="#00247D").grid(row=0, column=0, columnspan=3, pady=10)

        ttk.Label(f, text="Topic Query:").grid(row=1, column=0, sticky="w", pady=4)
        self.e_topic = ttk.Entry(f, width=40)
        self.e_topic.grid(row=1, column=1, sticky="ew", pady=4, padx=5)
        self.e_topic.insert(0, "blueprint modern digital government")
        ttk.Button(f, text="⭐ Save Fav Topic", command=self.add_favorite_topic).grid(row=1, column=2, padx=5)

        self.exact_match_var = tk.BooleanVar(value=False)
        ttk.Checkbutton(f, text="Exact Phrase Match Only", variable=self.exact_match_var).grid(
            row=2, column=1, sticky="w", padx=5)

        ttk.Label(f, text="Target Count:").grid(row=3, column=0, sticky="w", pady=4)
        self.sp_count = ttk.Spinbox(f, from_=10, to=200, increment=10, width=10)
        self.sp_count.grid(row=3, column=1, sticky="w", pady=4, padx=5)
        self.sp_count.set(20)

        ttk.Label(f, text="Sort By:").grid(row=4, column=0, sticky="w", pady=4)
        self.cb_sort = ttk.Combobox(f, values=["Best Match", "Most Recent"], state="readonly", width=15)
        self.cb_sort.grid(row=4, column=1, sticky="w", pady=4, padx=5)
        self.cb_sort.set("Best Match")

        ttk.Label(f, text="Department:").grid(row=5, column=0, sticky="w", pady=4)
        self.cb_dept = ttk.Combobox(f, values=["All Departments", *config.DEPARTMENT_SLUGS.keys()],
                                    state="readonly", width=25)
        self.cb_dept.grid(row=5, column=1, sticky="w", pady=4, padx=5)
        self.cb_dept.set("All Departments")

        ttk.Label(f, text="Document Type:").grid(row=6, column=0, sticky="w", pady=4)
        self.cb_doc_type = ttk.Combobox(f, values=["All Types", *config.DOC_TYPE_SLUGS.keys()],
                                        state="readonly", width=25)
        self.cb_doc_type.grid(row=6, column=1, sticky="w", pady=4, padx=5)
        self.cb_doc_type.set("All Types")

        fav_frame = ttk.LabelFrame(f, text=" ⭐ Permanent Topic Favorites (Click to Load) ", padding=8)
        fav_frame.grid(row=7, column=0, columnspan=3, sticky="ew", pady=8)
        self.fav_buttons_inner = ttk.Frame(fav_frame)
        self.fav_buttons_inner.pack(fill="x")
        self.render_favorite_topics()

        btn_frame = ttk.Frame(f)
        btn_frame.grid(row=8, column=0, columnspan=3, pady=10)
        ttk.Button(btn_frame, text="🚀 Run Deep Search", command=self.start_harvest).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="⚔️ Split-Screen Compare", command=self.open_split_comparator).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="🔍 Cross-PDF Search", command=self.cross_pdf_search_dialog).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="📜 Search History", command=self.open_history_dialog).pack(side="left", padx=4)
        ttk.Button(btn_frame, text="📄 Export Word Brief (.docx)", command=self.export_word_briefing).pack(side="left", padx=4)

        self.log_box = tk.Text(f, height=6, bg="#1e293b", fg="#f8fafc", font=("Consolas", 10))
        self.log_box.grid(row=10, column=0, columnspan=3, sticky="nsew", pady=10)

    def render_favorite_topics(self) -> None:
        for widget in self.fav_buttons_inner.winfo_children():
            widget.destroy()
        if not self.state.favorite_topics:
            ttk.Label(self.fav_buttons_inner, text="No favorite search terms saved yet.",
                      font=("Segoe UI", 9, "italic")).pack()
            return
        for fav in self.state.favorite_topics:
            btn_f = ttk.Frame(self.fav_buttons_inner)
            btn_f.pack(side="left", padx=4, pady=2)
            ttk.Button(btn_f, text=f"⭐ {fav}", command=lambda t=fav: self.load_fav_topic_search(t)).pack(side="left")
            ttk.Button(btn_f, text="❌", width=2, command=lambda t=fav: self.remove_favorite_topic(t)).pack(side="left", padx=2)

    def add_favorite_topic(self) -> None:
        if self.state.add_favorite_topic(self.e_topic.get().strip()):
            self.render_favorite_topics()

    def remove_favorite_topic(self, topic: str) -> None:
        if self.state.remove_favorite_topic(topic):
            self.render_favorite_topics()

    def load_fav_topic_search(self, topic: str) -> None:
        self.e_topic.delete(0, tk.END)
        self.e_topic.insert(0, topic)
        self.start_harvest()

    def log(self, msg: str) -> None:
        self.root.after(0, lambda: self._log_on_main_thread(msg))

    def _log_on_main_thread(self, msg: str) -> None:
        self.log_box.insert(tk.END, msg + "\n")
        self.log_box.see(tk.END)

    def start_harvest(self) -> None:
        topic = self.e_topic.get().strip()
        if not topic:
            return
        self.state.record_search(topic, self.cb_dept.get(), self.cb_doc_type.get())
        threading.Thread(target=self._run_harvest, args=(topic,), daemon=True).start()

    def _run_harvest(self, topic: str) -> None:
        self.log(f"🔎 Scanning GOV.UK for '{topic}'...")
        try:
            target_total = int(self.sp_count.get())
        except (tk.TclError, ValueError):
            target_total = 20

        raw_results = deep_harvest_gov_uk(
            topic, target_total, self.cb_sort.get(), self.cb_dept.get(),
            self.cb_doc_type.get(), self.exact_match_var.get(), log_cb=self.log,
        )
        topic_dir, docs, _suggestions = archive.build_and_save_archive(topic, raw_results, log_cb=self.log)
        if docs:
            self.log(f"✅ Saved {len(docs)} documents for '{topic}'.")
            self.root.after(0, lambda: self._sync_reader(topic_dir, docs))

    def _sync_reader(self, topic_dir: Path, docs: list[Document]) -> None:
        self.active_docs = docs
        briefing_text = archive.load_briefing_text(topic_dir)
        self.txt_briefing.delete("1.0", tk.END)
        self.txt_briefing.insert(tk.END, briefing_text)
        self.refresh_doc_list()
        self.generate_analytics_matrix()
        self.nb.select(self.tab_reader)

    # ======================================================================
    # Reader tab
    # ======================================================================
    def _build_reader_tab(self) -> None:
        main_paned = ttk.Panedwindow(self.tab_reader, orient="horizontal")
        main_paned.pack(fill="both", expand=True, padx=5, pady=5)

        left = ttk.LabelFrame(main_paned, text=" Discovered Sources ", padding=8)
        main_paned.add(left, weight=1)

        self.e_filter = ttk.Entry(left)
        self.e_filter.pack(fill="x", pady=4)
        self.e_filter.bind("<KeyRelease>", lambda e: self.refresh_doc_list())

        self.lb_docs = tk.Listbox(left, bg="#ffffff", selectbackground="#00247D", exportselection=False)
        self.lb_docs.pack(fill="both", expand=True)
        self.lb_docs.bind("<<ListboxSelect>>", self.on_doc_select)
        self.lb_docs.bind("<Double-1>", self.on_doc_double_click)

        right_v_paned = ttk.Panedwindow(main_paned, orient="vertical")
        main_paned.add(right_v_paned, weight=4)

        right_top_frame = ttk.LabelFrame(right_v_paned, text=" Document Metadata & Attachments ", padding=6)
        right_v_paned.add(right_top_frame, weight=1)

        btn_row = ttk.Frame(right_top_frame)
        btn_row.pack(fill="x", pady=(0, 4))
        self.btn_star = ttk.Button(btn_row, text="⭐ Star Source", command=self.toggle_starred)
        self.btn_star.pack(side="left", padx=2)
        ttk.Button(btn_row, text="🏷️ Set Tag", command=self.apply_custom_doc_tag).pack(side="left", padx=2)
        ttk.Button(btn_row, text="📋 Copy Citation", command=self.copy_citation_popup).pack(side="left", padx=2)

        self.reader_details_att_paned = ttk.Panedwindow(right_top_frame, orient="horizontal")
        self.reader_details_att_paned.pack(fill="both", expand=True)

        details_box_frame = ttk.Frame(self.reader_details_att_paned)
        self.reader_details_att_paned.add(details_box_frame, weight=1)
        self.txt_details = scrolledtext.ScrolledText(details_box_frame, height=4, wrap="word")
        self.txt_details.pack(fill="both", expand=True)

        att_box_frame = ttk.LabelFrame(self.reader_details_att_paned, text=" Attachments ", padding=4)
        self.reader_details_att_paned.add(att_box_frame, weight=1)

        sb_att = ttk.Scrollbar(att_box_frame, orient="vertical")
        self.lb_atts = tk.Listbox(att_box_frame, height=4, exportselection=False, yscrollcommand=sb_att.set)
        sb_att.config(command=self.lb_atts.yview)
        sb_att.pack(side="right", fill="y")
        self.lb_atts.pack(side="left", fill="both", expand=True)
        self.lb_atts.bind("<Double-1>", self.on_attachment_open)

        def _set_reader_sash(event):
            w = event.width
            if w > 100:
                self.reader_details_att_paned.sashpos(0, int(w * 0.5))
                self.reader_details_att_paned.unbind("<Configure>")

        self.reader_details_att_paned.bind("<Configure>", _set_reader_sash)

        right_bottom_frame = ttk.Frame(right_v_paned)
        right_v_paned.add(right_bottom_frame, weight=5)

        viewer_hdr = ttk.Frame(right_bottom_frame)
        viewer_hdr.pack(fill="x", pady=(0, 2))
        ttk.Button(viewer_hdr, text="🔍 Pop Out Viewer Window", command=self.pop_out_reader_pdf).pack(side="right", padx=2)

        self.reader_nb = ttk.Notebook(right_bottom_frame)
        self.reader_nb.pack(fill="both", expand=True)

        self.reader_pdf_viewer = PDFViewerWidget(self.reader_nb, lambda: self.state.keyword_rules)
        self.reader_nb.add(self.reader_pdf_viewer, text="📕 PDF Viewer")

        briefing_frame = ttk.Frame(self.reader_nb)
        self.reader_nb.add(briefing_frame, text="📰 Briefing")

        briefing_tb = ttk.Frame(briefing_frame, padding=4)
        briefing_tb.pack(fill="x", side="top")

        ttk.Button(
            briefing_tb,
            text="✨ Generate AI Briefing",
            command=self.run_ai_briefing
        ).pack(side="left", padx=4)

        self.lbl_ai_status = ttk.Label(briefing_tb, text="", font=("Segoe UI", 9, "italic"))
        self.lbl_ai_status.pack(side="left", padx=8)
        
        # Check local Ollama readiness in background
        threading.Thread(target=self.check_ai_readiness, daemon=True).start()

        self.txt_briefing = scrolledtext.ScrolledText(briefing_frame, wrap="word")
        self.txt_briefing.pack(fill="both", expand=True)

    def check_ai_readiness(self) -> None:
        """Checks if Ollama is live and updates the UI status badge."""
        if ai_engine.check_ollama_status():
            self.root.after(0, lambda: self.lbl_ai_status.config(text="🟢 Local AI (Ollama) Ready", foreground="#16a34a"))
        else:
            self.root.after(0, lambda: self.lbl_ai_status.config(text="⚡ Using Local NLP Fallback (Ollama offline)", foreground="#d97706"))

    def run_ai_briefing(self) -> None:
        current_text = ""
        if self.reader_pdf_viewer.doc_obj:
            full_pdf_text = []
            for page in self.reader_pdf_viewer.doc_obj:
                textpage = page.get_textpage()
                full_pdf_text.append(textpage.get_text_range())
            current_text = "\n".join(full_pdf_text)

        if not current_text.strip() and self.selected_doc:
            current_text = f"{self.selected_doc.title}\n{self.selected_doc.description}"
        if not current_text.strip():
            current_text = self.txt_briefing.get("1.0", tk.END).strip()

        if not current_text:
            messagebox.showwarning("No Text Loaded", "Please select a document or open a PDF attachment first.")
            return

        self.lbl_ai_status.config(text="🤖 Analyzing document with local AI...")

        def _worker():
            summary = ai_engine.generate_document_briefing(current_text)
            self.root.after(0, lambda: self._display_ai_summary(summary))

        threading.Thread(target=_worker, daemon=True).start()

    def _display_ai_summary(self, summary_text: str) -> None:
        self.txt_briefing.delete("1.0", tk.END)
        self.txt_briefing.insert(tk.END, summary_text)
        self.check_ai_readiness()

    def refresh_doc_list(self) -> None:
        self.lb_docs.delete(0, tk.END)
        q = self.e_filter.get().strip().lower()
        for doc in self.active_docs:
            if q and q not in doc.title.lower() and q not in doc.description.lower():
                continue
            icon = "⭐" if self.state.is_favorite(doc) else "📄"
            tag = self.state.get_tag(doc)
            tag_label = f" [{tag}]" if tag else ""
            self.lb_docs.insert(tk.END, f"{icon} {doc.title}{tag_label}")

    def _visible_docs(self) -> list[Document]:
        q = self.e_filter.get().strip().lower()
        if not q:
            return self.active_docs
        return [d for d in self.active_docs if q in d.title.lower() or q in d.description.lower()]

    def on_doc_select(self, _event) -> None:
        sel = self.lb_docs.curselection()
        if not sel:
            return
        docs = self._visible_docs()
        if sel[0] >= len(docs):
            return
        doc = docs[sel[0]]
        self.selected_doc = doc

        details = (
            f"SOURCE: {doc.title}\nDATE: {doc.date}\nURL: {doc.url}\n"
            f"{'-' * 40}\n{doc.description}"
        )
        self.txt_details.delete("1.0", tk.END)
        self.txt_details.insert(tk.END, details)

        self.btn_star.config(text="🌟 Unstar Source" if self.state.is_favorite(doc) else "⭐ Star Source")

        self.active_atts = doc.attachments
        self.lb_atts.delete(0, tk.END)
        for url in self.active_atts:
            self.lb_atts.insert(tk.END, _attachment_label(url))

    def on_doc_double_click(self, _event) -> None:
        if self.selected_doc:
            webbrowser.open_new_tab(self.selected_doc.url)

    def toggle_starred(self) -> None:
        if not self.selected_doc:
            return
        is_fav = self.state.toggle_favorite_source(self.selected_doc)
        self.btn_star.config(text="🌟 Unstar Source" if is_fav else "⭐ Star Source")
        self.refresh_doc_list()
        self.refresh_fav_hub()

    def apply_custom_doc_tag(self) -> None:
        if not self.selected_doc:
            return
        tag = simpledialog.askstring("Set Custom Tag", "Enter custom label for this document (e.g., 'High Priority', 'Legal'):")
        if tag:
            self.state.set_tag(self.selected_doc, tag)
            self.refresh_doc_list()

    def on_attachment_open(self, _event) -> None:
        sel = self.lb_atts.curselection()
        if sel and self.active_atts:
            self._route_attachment_open(self.active_atts[sel[0]], viewer=self.reader_pdf_viewer, notebook=self.reader_nb)

    def pop_out_reader_pdf(self) -> None:
        if not self.reader_pdf_viewer.pdf_path:
            messagebox.showinfo("No Document Loaded", "Open a PDF attachment first to pop out the viewer.")
            return
        pop = tk.Toplevel(self.root)
        pop.title(f"📖 {os.path.basename(self.reader_pdf_viewer.pdf_path)}")
        pop.geometry("1100x850")
        pv = PDFViewerWidget(pop, lambda: self.state.keyword_rules)
        pv.pack(fill="both", expand=True)
        pv.load_pdf(self.reader_pdf_viewer.pdf_path)

        if self.reader_pdf_viewer.auto_highlights_index:
            pv.auto_highlight(show_dialog=False)

    def pop_out_fav_pdf(self) -> None:
        if not self.fav_pdf_viewer.pdf_path:
            messagebox.showinfo("No Document Loaded", "Open a PDF attachment first to pop out the viewer.")
            return
        pop = tk.Toplevel(self.root)
        pop.title(f"📖 {os.path.basename(self.fav_pdf_viewer.pdf_path)}")
        pop.geometry("1100x850")
        pv = PDFViewerWidget(pop, lambda: self.state.keyword_rules)
        pv.pack(fill="both", expand=True)
        pv.load_pdf(self.fav_pdf_viewer.pdf_path)

        if self.fav_pdf_viewer.auto_highlights_index:
            pv.auto_highlight(show_dialog=False)

    def _route_attachment_open(self, url: str, viewer: PDFViewerWidget, notebook: ttk.Notebook) -> None:
        tag = classify_attachment_url(url)
        if "📕 PDF" in tag:
            self._open_pdf_attachment(url, viewer, notebook)
        elif "📊 Data" in tag:
            DataGridViewerWindow(self.root, url, title=url.rsplit("/", 1)[-1])
        else:
            webbrowser.open_new_tab(url)

    def _open_pdf_attachment(
        self,
        url: str,
        viewer: PDFViewerWidget,
        notebook: ttk.Notebook,
        on_loaded: Callable[[], None] | None = None,
    ) -> None:
        pdf_name = sanitize_filename(url)
        local_p = os.path.abspath(os.path.join(config.DATA_DIR, pdf_name))

        def _finish():
            viewer.load_pdf(local_p)
            notebook.select(viewer)
            if on_loaded:
                on_loaded()

        if os.path.exists(local_p):
            _finish()
            return

        self.log(f"⬇️ Downloading {pdf_name}...")

        def _download():
            try:
                res = requests.get(url, timeout=30)
                res.raise_for_status()
            except requests.RequestException as exc:
                logger.warning("Failed to download PDF %s: %s", url, exc)
                self.root.after(0, lambda: messagebox.showerror("Download Failed", f"Could not download PDF:\n{exc}"))
                return
            try:
                os.makedirs(os.path.dirname(local_p), exist_ok=True)
                with open(local_p, "wb") as f:
                    f.write(res.content)
            except OSError as exc:
                self.root.after(0, lambda: messagebox.showerror("Save Failed", f"Could not save PDF locally:\n{exc}"))
                return
            self.root.after(0, _finish)

        threading.Thread(target=_download, daemon=True).start()

    # ======================================================================
    # Favorites Hub tab
    # ======================================================================
    def _build_favs_tab(self) -> None:
        main_paned = ttk.Panedwindow(self.tab_favs, orient="horizontal")
        main_paned.pack(fill="both", expand=True, padx=5, pady=5)

        left = ttk.LabelFrame(main_paned, text=" ⭐ Starred Sources ", padding=8)
        main_paned.add(left, weight=1)
        self.lb_favs = tk.Listbox(left, exportselection=False)
        self.lb_favs.pack(fill="both", expand=True)
        self.lb_favs.bind("<<ListboxSelect>>", self.on_fav_select)
        self.lb_favs.bind("<Double-1>", self.on_fav_double_click)

        right_v_paned = ttk.Panedwindow(main_paned, orient="vertical")
        main_paned.add(right_v_paned, weight=4)

        right_top = ttk.LabelFrame(right_v_paned, text=" Source Metadata & Assets ", padding=6)
        right_v_paned.add(right_top, weight=1)

        self.fav_det_att_paned = ttk.Panedwindow(right_top, orient="horizontal")
        self.fav_det_att_paned.pack(fill="both", expand=True)

        fav_det_frame = ttk.Frame(self.fav_det_att_paned)
        self.fav_det_att_paned.add(fav_det_frame, weight=1)
        self.txt_fav_details = scrolledtext.ScrolledText(fav_det_frame, height=4, wrap="word")
        self.txt_fav_details.pack(fill="both", expand=True)

        fav_att_frame = ttk.LabelFrame(self.fav_det_att_paned, text=" Attachments ", padding=4)
        self.fav_det_att_paned.add(fav_att_frame, weight=1)

        sb_fav_att = ttk.Scrollbar(fav_att_frame, orient="vertical")
        self.lb_fav_atts = tk.Listbox(fav_att_frame, height=4, exportselection=False, yscrollcommand=sb_fav_att.set)
        sb_fav_att.config(command=self.lb_fav_atts.yview)
        sb_fav_att.pack(side="right", fill="y")
        self.lb_fav_atts.pack(side="left", fill="both", expand=True)
        self.lb_fav_atts.bind("<Double-1>", self.on_fav_att_open)

        def _set_fav_sash(event):
            w = event.width
            if w > 100:
                self.fav_det_att_paned.sashpos(0, int(w * 0.5))
                self.fav_det_att_paned.unbind("<Configure>")

        self.fav_det_att_paned.bind("<Configure>", _set_fav_sash)

        right_bottom = ttk.Frame(right_v_paned)
        right_v_paned.add(right_bottom, weight=5)

        fav_hdr = ttk.Frame(right_bottom)
        fav_hdr.pack(fill="x", pady=(0, 2))
        ttk.Button(fav_hdr, text="🔍 Pop Out Viewer Window", command=self.pop_out_fav_pdf).pack(side="right", padx=2)

        self.fav_nb = ttk.Notebook(right_bottom)
        self.fav_nb.pack(fill="both", expand=True)
        self.fav_pdf_viewer = PDFViewerWidget(self.fav_nb, lambda: self.state.keyword_rules)
        self.fav_nb.add(self.fav_pdf_viewer, text="📕 PDF Preview")

    def _favorite_ids(self) -> list[str]:
        return list(self.state.favorite_sources.keys())

    def refresh_fav_hub(self) -> None:
        self.lb_favs.delete(0, tk.END)
        self.all_fav_attachments = {"pdf": [], "link": [], "data": []}

        for doc_id in self._favorite_ids():
            entry = self.state.favorite_sources[doc_id]
            self.lb_favs.insert(tk.END, f"⭐ {entry['title']}")
            for url in entry.get("attachments", []):
                tag = classify_attachment_url(url)
                bucket = "pdf" if "📕 PDF" in tag else ("data" if "📊 Data" in tag else "link")
                self.all_fav_attachments[bucket].append(url)

    def on_fav_select(self, _event) -> None:
        ids = self._favorite_ids()
        sel = self.lb_favs.curselection()
        if not sel:
            return
        doc_id = ids[sel[0]]
        entry = self.state.favorite_sources[doc_id]
        self.fav_selected_doc_id = doc_id
        self.fav_active_atts = list(entry.get("attachments", []))

        self.txt_fav_details.delete("1.0", tk.END)
        self.txt_fav_details.insert(
            tk.END, f"TITLE: {entry['title']}\nURL: {entry['url']}\nTOPIC: {entry.get('topic', '')}"
        )

        self.lb_fav_atts.delete(0, tk.END)
        for url in self.fav_active_atts:
            self.lb_fav_atts.insert(tk.END, _attachment_label(url))

    def on_fav_double_click(self, _event) -> None:
        ids = self._favorite_ids()
        sel = self.lb_favs.curselection()
        if sel:
            webbrowser.open_new_tab(self.state.favorite_sources[ids[sel[0]]]["url"])

    def on_fav_att_open(self, _event) -> None:
        sel = self.lb_favs.curselection()
        if sel and self.fav_active_atts:
            self._route_attachment_open(self.fav_active_atts[sel[0]], viewer=self.fav_pdf_viewer, notebook=self.fav_nb)

    # ======================================================================
    # Keyword Brain tab
    # ======================================================================
    def _build_kw_tab(self) -> None:
        f = ttk.Frame(self.tab_kw, padding=12)
        f.pack(fill="both", expand=True)

        overview_frame = ttk.LabelFrame(
            f, text=" 🌐 Master Overview — All Active Categories (click a card to jump to it) ", padding=10
        )
        overview_frame.pack(fill="x", side="top", pady=(0, 10))

        overview_hdr = ttk.Frame(overview_frame)
        overview_hdr.pack(fill="x", pady=(0, 6))
        self.lbl_overview_summary = ttk.Label(
            overview_hdr, text="", font=("Segoe UI", 9, "italic"), foreground="#475569"
        )
        self.lbl_overview_summary.pack(side="left")
        ttk.Button(
            overview_hdr, text="🌐 Load Full Master Overview Profile", command=self.load_master_overview_profile
        ).pack(side="right")

        self.master_overview_container = ttk.Frame(overview_frame)
        self.master_overview_container.pack(fill="x")

        preset_bar = ttk.LabelFrame(f, text=" 📦 Research Profiles & Presets ", padding=8)
        preset_bar.pack(fill="x", side="top", pady=(0, 10))

        ttk.Label(preset_bar, text="Profile:").pack(side="left", padx=(4, 4))
        self.cb_kw_presets = ttk.Combobox(
            preset_bar,
            values=list(self.state.research_profiles.keys()),
            state="readonly",
            width=36,
        )
        self.cb_kw_presets.pack(side="left", padx=4)

        ttk.Button(preset_bar, text="📥 Load Profile", command=self.load_kw_preset).pack(side="left", padx=4)
        ttk.Button(preset_bar, text="💾 Save Profile As...", command=self.save_custom_profile).pack(side="left", padx=4)
        ttk.Button(preset_bar, text="🗑️ Delete Profile", command=self.delete_custom_profile).pack(side="left", padx=4)
        ttk.Button(preset_bar, text="📤 Export to File...", command=self.export_profile_to_file).pack(side="left", padx=(16, 4))
        ttk.Button(preset_bar, text="📂 Import from File...", command=self.import_profile_from_file).pack(side="left", padx=4)

        self._refresh_profile_combobox()

        body = ttk.Frame(f)
        body.pack(fill="both", expand=True)

        left = ttk.LabelFrame(body, text=" Categories & Highlight Palette ", padding=10)
        left.pack(side="left", fill="both", expand=False, padx=(0, 8))

        self.tree_kw_cats = ttk.Treeview(
            left, columns=("swatch", "name", "count"), show="headings", height=12, selectmode="browse"
        )
        self.tree_kw_cats.heading("swatch", text="Color", anchor="center")
        self.tree_kw_cats.heading("name", text="Category Name", anchor="w")
        self.tree_kw_cats.heading("count", text="Terms", anchor="center")

        self.tree_kw_cats.column("swatch", width=55, anchor="center")
        self.tree_kw_cats.column("name", width=190, anchor="w")
        self.tree_kw_cats.column("count", width=65, anchor="center")

        sb_cats = ttk.Scrollbar(left, orient="vertical", command=self.tree_kw_cats.yview)
        self.tree_kw_cats.configure(yscrollcommand=sb_cats.set)

        sb_cats.pack(side="right", fill="y")
        self.tree_kw_cats.pack(side="top", fill="both", expand=True)
        self.tree_kw_cats.bind("<<TreeviewSelect>>", self.on_kw_cat_select)

        self.color_preview_frame = ttk.Frame(left, padding=6)
        self.color_preview_frame.pack(fill="x", side="top", pady=(8, 4))

        self.lbl_selected_cat_name = ttk.Label(
            self.color_preview_frame, text="Active: (Select a category)", font=("Segoe UI", 9, "bold")
        )
        self.lbl_selected_cat_name.pack(side="left", padx=(2, 6))

        self.canvas_color_swatch = tk.Canvas(self.color_preview_frame, width=28, height=18, bg="#ffef33", highlightthickness=1)
        self.canvas_color_swatch.pack(side="left", padx=2)

        color_box = ttk.LabelFrame(left, text=" 🎨 Assign Highlight Color ", padding=6)
        color_box.pack(fill="x", side="top", pady=(4, 6))

        palette_frame = ttk.Frame(color_box)
        palette_frame.pack(fill="x")

        for idx, (label, rgb, hex_val) in enumerate(PALETTE_PRESETS):
            r_idx, c_idx = divmod(idx, 4)
            btn = tk.Button(
                palette_frame,
                text=label,
                bg=hex_val,
                fg="#000000" if label in ["Yellow", "Green", "Pink"] else "#ffffff",
                font=("Segoe UI", 8, "bold"),
                relief="groove",
                command=lambda color_val=rgb: self.set_selected_category_color(color_val)
            )
            btn.grid(row=r_idx, column=c_idx, padx=2, pady=2, sticky="ew")

        palette_frame.columnconfigure((0, 1, 2, 3), weight=1)

        ttk.Button(color_box, text="🎨 Custom Color...", command=self.pick_custom_category_color).pack(fill="x", pady=(4, 0))

        cat_ctrl = ttk.Frame(left)
        cat_ctrl.pack(fill="x", side="top", pady=(6, 0))

        self.e_new_cat = ttk.Entry(cat_ctrl)
        self.e_new_cat.pack(fill="x", pady=(0, 4))

        btn_row = ttk.Frame(cat_ctrl)
        btn_row.pack(fill="x")
        ttk.Button(btn_row, text="➕ Add Category", command=self.add_kw_category).pack(side="left", expand=True, fill="x", padx=(0, 2))
        ttk.Button(btn_row, text="🗑️ Remove Category", command=self.remove_kw_category).pack(side="left", expand=True, fill="x", padx=(2, 0))

        right = ttk.LabelFrame(body, text=" Associated Research Keywords ", padding=10)
        right.pack(side="left", fill="both", expand=True)

        self.lbl_terms_header = ttk.Label(
            right, text="Double-click any keyword to toggle active (☑) or inactive (☐)", font=("Segoe UI", 9, "italic")
        )
        self.lbl_terms_header.pack(anchor="w", pady=(0, 6))

        self.lb_kw_terms = tk.Listbox(right, exportselection=False, font=("Segoe UI", 10), bg="#ffffff")
        sb_terms = ttk.Scrollbar(right, orient="vertical", command=self.lb_kw_terms.yview)
        self.lb_kw_terms.configure(yscrollcommand=sb_terms.set)

        sb_terms.pack(side="right", fill="y")
        self.lb_kw_terms.pack(side="top", fill="both", expand=True)
        self.lb_kw_terms.bind("<Double-1>", self.toggle_kw_term_state)

        term_ctrl = ttk.Frame(right)
        term_ctrl.pack(fill="x", pady=(8, 0))
        self.e_new_kw = ttk.Entry(term_ctrl)
        self.e_new_kw.pack(side="left", fill="x", expand=True, padx=(0, 4))
        ttk.Button(term_ctrl, text="➕ Add Keyword", command=self.add_kw_term).pack(side="left", padx=2)
        ttk.Button(term_ctrl, text="🗑️ Remove Keyword", command=self.remove_kw_term).pack(side="left", padx=2)

        self.refresh_kw_categories_list()

    def _refresh_profile_combobox(self) -> None:
        profiles = list(self.state.research_profiles.keys())
        self.cb_kw_presets.config(values=profiles)
        if profiles:
            self.cb_kw_presets.set(profiles[0])

    def load_kw_preset(self) -> None:
        preset_key = self.cb_kw_presets.get()
        if self.state.apply_preset_template(preset_key):
            self.refresh_kw_categories_list()
            self.generate_analytics_matrix()
            messagebox.showinfo("Profile Loaded", f"Loaded research profile '{preset_key}'!")

    def save_custom_profile(self) -> None:
        name = simpledialog.askstring("Save Research Profile", "Enter a name for this custom profile:")
        if name:
            if self.state.save_current_as_profile(name):
                self._refresh_profile_combobox()
                self.cb_kw_presets.set(name)
                messagebox.showinfo("Profile Saved", f"Saved profile '{name}'!")

    def delete_custom_profile(self) -> None:
        selected = self.cb_kw_presets.get()
        if not selected:
            return
        if messagebox.askyesno("Delete Profile", f"Are you sure you want to delete profile '{selected}'?"):
            if self.state.delete_profile(selected):
                self._refresh_profile_combobox()
                messagebox.showinfo("Profile Deleted", f"Deleted profile '{selected}'.")

    def export_profile_to_file(self) -> None:
        """Writes the currently-active keyword rules to a shareable JSON file.

        Free/local by design: no account, no cloud sync -- just a plain JSON
        file that another user of this app can drop straight back in.
        """
        selected = self.cb_kw_presets.get() or "Custom Profile"
        out_p = filedialog.asksaveasfilename(
            defaultextension=".json",
            filetypes=[("Keyword Brain Profile", "*.json")],
            initialfile=f"{selected}.json".replace(" ", "_"),
        )
        if not out_p:
            return
        payload = {"profile_name": selected, "rules": self.state.keyword_rules}
        try:
            with open(out_p, "w", encoding="utf-8") as fh:
                json.dump(payload, fh, indent=2, ensure_ascii=False)
        except OSError as exc:
            messagebox.showerror("Export Failed", f"Could not write file:\n{exc}")
            return
        messagebox.showinfo("Profile Exported", f"Saved to: {out_p}")

    def import_profile_from_file(self) -> None:
        """Loads a keyword-rule profile JSON file exported by this app (or hand-written)."""
        in_p = filedialog.askopenfilename(filetypes=[("Keyword Brain Profile", "*.json"), ("All files", "*.*")])
        if not in_p:
            return
        try:
            with open(in_p, "r", encoding="utf-8") as fh:
                data = json.load(fh)
        except (OSError, json.JSONDecodeError) as exc:
            messagebox.showerror("Import Failed", f"Could not read file:\n{exc}")
            return

        if isinstance(data, dict) and "rules" in data:
            profile_name = str(data.get("profile_name") or Path(in_p).stem)
            rules = data["rules"]
        else:
            # Allow importing a bare {"category": {...}} rules dict too.
            profile_name = Path(in_p).stem
            rules = data

        if not isinstance(rules, dict):
            messagebox.showerror("Import Failed", "That file doesn't look like a valid keyword profile.")
            return

        self.state.research_profiles[profile_name] = normalize_keyword_rules(rules)
        save_json(PROFILES_FILE, self.state.research_profiles)

        self._refresh_profile_combobox()
        self.cb_kw_presets.set(profile_name)
        messagebox.showinfo("Profile Imported", f"Imported profile '{profile_name}'. Click 'Load Profile' to activate it.")

    def refresh_kw_categories_list(self) -> None:
        for row in self.tree_kw_cats.get_children():
            self.tree_kw_cats.delete(row)

        for cat, data in self.state.keyword_rules.items():
            rgb = data.get("color", [1.0, 0.9, 0.2])
            hex_val = _rgb_to_hex(rgb)
            terms_count = len(data.get("terms", {}))

            item_id = self.tree_kw_cats.insert(
                "", "end", iid=cat, values=("███", cat, f"{terms_count} terms")
            )
            self.tree_kw_cats.tag_configure(f"tag_{cat}", foreground=hex_val)
            self.tree_kw_cats.item(item_id, tags=(f"tag_{cat}",))

        children = self.tree_kw_cats.get_children()
        if children:
            self.tree_kw_cats.selection_set(children[0])
            self.on_kw_cat_select(None)

        self.refresh_master_overview()

    def load_master_overview_profile(self, silent: bool = False) -> None:
        """One-click shortcut: activate the auto-aggregated profile containing every
        category from every saved research profile, so the card grid shows everything."""
        master_key = "🌐 Master Intelligence Overview"
        if master_key not in self.state.research_profiles:
            if not silent:
                messagebox.showinfo(
                    "No Profiles Yet",
                    "No saved research profiles were found to aggregate. Add categories below, "
                    "or save/import a profile first.",
                )
            return
        self.state.apply_preset_template(master_key)
        self.cb_kw_presets.set(master_key)
        self.refresh_kw_categories_list()
        self.generate_analytics_matrix()

    def refresh_master_overview(self) -> None:
        """Renders a scannable card grid summarizing every category currently active,
        so the user gets the big picture before drilling into individual term lists."""
        for w in self.master_overview_container.winfo_children():
            w.destroy()

        categories = list(self.state.keyword_rules.items())
        if not categories:
            self.lbl_overview_summary.config(text="No categories yet — add one below to get started.")
            return

        total_terms = sum(len(data.get("terms", {})) for _cat, data in categories)
        self.lbl_overview_summary.config(
            text=f"{len(categories)} categories  •  {total_terms} keyword terms total"
        )

        cols = 4
        for idx, (cat, data) in enumerate(categories):
            row, col = divmod(idx, cols)
            terms = data.get("terms", {})
            active_n = sum(1 for enabled in terms.values() if enabled)
            total_n = len(terms)
            hex_val = _rgb_to_hex(data.get("color", [1.0, 0.9, 0.2]))

            stats_text = f"{active_n}/{total_n} terms active"
            if self.active_docs:
                hit_total = 0
                for doc in self.active_docs:
                    haystack = f"{doc.title} {doc.description}".lower()
                    hit_total += self._category_hit_count(haystack, cat)
                stats_text += f"  •  {hit_total} hits in loaded docs"

            card = tk.Frame(
                self.master_overview_container, bg="#ffffff",
                highlightbackground="#cbd5e1", highlightthickness=1, cursor="hand2",
            )
            card.grid(row=row, column=col, padx=5, pady=5, sticky="nsew")

            accent = tk.Frame(card, bg=hex_val, width=6)
            accent.pack(side="left", fill="y")

            body = tk.Frame(card, bg="#ffffff", padx=10, pady=8)
            body.pack(side="left", fill="both", expand=True)

            name_lbl = tk.Label(
                body, text=cat, font=("Segoe UI", 10, "bold"), bg="#ffffff", fg="#0f172a",
                anchor="w", wraplength=200, justify="left",
            )
            name_lbl.pack(anchor="w", fill="x")

            stats_lbl = tk.Label(
                body, text=stats_text, font=("Segoe UI", 8), bg="#ffffff", fg="#64748b", anchor="w"
            )
            stats_lbl.pack(anchor="w", fill="x")

            for widget in (card, accent, body, name_lbl, stats_lbl):
                widget.bind("<Button-1>", lambda _e, c=cat: self.select_kw_category_from_overview(c))

        # Trailing "+ Add Category" card so a custom category can be created
        # in one click straight from the overview, without hunting for the
        # smaller form under the tree list below.
        add_idx = len(categories)
        row, col = divmod(add_idx, cols)
        add_card = tk.Frame(
            self.master_overview_container, bg="#f8fafc",
            highlightbackground="#94a3b8", highlightthickness=1, cursor="hand2",
        )
        add_card.grid(row=row, column=col, padx=5, pady=5, sticky="nsew")
        add_lbl = tk.Label(
            add_card, text="➕ Add New Category", font=("Segoe UI", 10, "bold"),
            bg="#f8fafc", fg="#334155", justify="center",
        )
        add_lbl.pack(fill="both", expand=True, padx=10, pady=18)
        for widget in (add_card, add_lbl):
            widget.bind("<Button-1>", lambda _e: self.add_category_from_overview())

        for c in range(cols):
            self.master_overview_container.columnconfigure(c, weight=1)

    def add_category_from_overview(self) -> None:
        """Quick-add flow: name the category, get an auto-assigned distinct color,
        and jump straight into its (empty) term list to start adding keywords."""
        name = simpledialog.askstring("Add New Category", "Enter a name for the new keyword category:")
        if not name or not name.strip():
            return
        name = name.strip()
        default_color = PALETTE_PRESETS[len(self.state.keyword_rules) % len(PALETTE_PRESETS)][1]
        if not self.state.add_keyword_category(name, default_color):
            messagebox.showwarning("Category Exists", f"A category named '{name}' already exists.")
            return
        self.refresh_kw_categories_list()
        self.generate_analytics_matrix()
        if name in self.tree_kw_cats.get_children():
            self.tree_kw_cats.selection_set(name)
            self.tree_kw_cats.see(name)
            self.on_kw_cat_select(None)

    def select_kw_category_from_overview(self, cat: str) -> None:
        if cat in self.tree_kw_cats.get_children():
            self.tree_kw_cats.selection_set(cat)
            self.tree_kw_cats.see(cat)
            self.on_kw_cat_select(None)

    def set_selected_category_color(self, rgb: list[float]) -> None:
        sel = self.tree_kw_cats.selection()
        if not sel:
            messagebox.showwarning("Select Category", "Please select a category from the list first.")
            return

        cat = sel[0]
        self.state.set_category_color(cat, rgb)
        self.refresh_kw_categories_list()
        self.tree_kw_cats.selection_set(cat)

    def pick_custom_category_color(self) -> None:
        sel = self.tree_kw_cats.selection()
        if not sel:
            messagebox.showwarning("Select Category", "Please select a category from the list first.")
            return

        cat = sel[0]
        curr_rgb = self.state.keyword_rules.get(cat, {}).get("color", [1.0, 0.9, 0.2])
        hex_col = _rgb_to_hex(curr_rgb)

        color_rgb, _color_hex = colorchooser.askcolor(initialcolor=hex_col, title=f"Choose Highlight Color for '{cat}'")
        if color_rgb:
            norm_color = [round(c / 255.0, 2) for c in color_rgb]
            self.state.set_category_color(cat, norm_color)
            self.refresh_kw_categories_list()
            self.tree_kw_cats.selection_set(cat)

    def add_kw_category(self) -> None:
        cat_name = self.e_new_cat.get().strip()
        if not cat_name:
            return
        if self.state.add_keyword_category(cat_name):
            self.e_new_cat.delete(0, tk.END)
            self.refresh_kw_categories_list()
            self.generate_analytics_matrix()

    def remove_kw_category(self) -> None:
        sel = self.tree_kw_cats.selection()
        if not sel:
            return
        cat = sel[0]
        if self.state.remove_keyword_category(cat):
            self.refresh_kw_categories_list()
            self.lb_kw_terms.delete(0, tk.END)
            self.generate_analytics_matrix()

    def on_kw_cat_select(self, _event) -> None:
        sel = self.tree_kw_cats.selection()
        if not sel:
            return
        cat = sel[0]
        data = self.state.keyword_rules.get(cat, {})
        curr_rgb = data.get("color", [1.0, 0.9, 0.2])
        hex_col = _rgb_to_hex(curr_rgb)

        self.lbl_selected_cat_name.config(text=f"Active: {cat}")
        self.canvas_color_swatch.config(bg=hex_col)

        self.lb_kw_terms.delete(0, tk.END)
        for term, enabled in data.get("terms", {}).items():
            self.lb_kw_terms.insert(tk.END, f"{'☑' if enabled else '☐'} {term}")

    def add_kw_term(self) -> None:
        sel = self.tree_kw_cats.selection()
        if not sel:
            return
        cat = sel[0]
        if self.state.add_keyword_term(cat, self.e_new_kw.get()):
            self.on_kw_cat_select(None)
            self.e_new_kw.delete(0, tk.END)
            self.refresh_kw_categories_list()
            self.tree_kw_cats.selection_set(cat)
            self.generate_analytics_matrix()

    def toggle_kw_term_state(self, event=None) -> None:
        cat_sel = self.tree_kw_cats.selection()
        if not cat_sel:
            return
        cat = cat_sel[0]

        index = self.lb_kw_terms.nearest(event.y) if (event and hasattr(event, "y")) else None
        if index is None:
            term_sel = self.lb_kw_terms.curselection()
            if not term_sel:
                return
            index = term_sel[0]

        item_str = self.lb_kw_terms.get(index)
        if not item_str:
            return
        term = item_str[2:].strip()
        self.state.toggle_keyword_term(cat, term)
        self.on_kw_cat_select(None)
        self.lb_kw_terms.select_set(index)
        self.generate_analytics_matrix()

    def remove_kw_term(self) -> None:
        cat_sel, term_sel = self.tree_kw_cats.selection(), self.lb_kw_terms.curselection()
        if not (cat_sel and term_sel):
            return
        cat = cat_sel[0]
        term = self.lb_kw_terms.get(term_sel[0])[2:].strip()
        if self.state.remove_keyword_term(cat, term):
            self.on_kw_cat_select(None)
            self.refresh_kw_categories_list()
            self.tree_kw_cats.selection_set(cat)
            self.generate_analytics_matrix()

    # ======================================================================
    # Analytics tab (Unified Scrollable Grid Matrix)
    # ======================================================================
    def _build_analytics_tab(self) -> None:
        f = ttk.Frame(self.tab_analytics, padding=10)
        f.pack(fill="both", expand=True)

        self.summary_card = ttk.LabelFrame(f, text=" 📊 Policy Intelligence Matrix ", padding=8)
        self.summary_card.pack(fill="x", side="top", pady=(0, 6))

        self.lbl_analytics_summary = ttk.Label(
            self.summary_card,
            text="Load sources via Control Panel to calculate category density matrix...",
            font=("Segoe UI", 10, "bold"),
            foreground="#00247D",
        )
        self.lbl_analytics_summary.pack(side="left", padx=4)

        triage_bar = ttk.LabelFrame(f, text=" 🔎 Triage: Sort & Filter ", padding=8)
        triage_bar.pack(fill="x", side="top", pady=(0, 6))

        ttk.Label(triage_bar, text="Sort by:").pack(side="left", padx=(4, 4))
        self.cb_analytics_sort = ttk.Combobox(
            triage_bar, values=["Relevance (Total Hits)", "Search Order"], state="readonly", width=22
        )
        self.cb_analytics_sort.pack(side="left", padx=4)
        self.cb_analytics_sort.set("Relevance (Total Hits)")
        self.cb_analytics_sort.bind("<<ComboboxSelected>>", lambda e: self.generate_analytics_matrix())

        ttk.Label(triage_bar, text="Focus category:").pack(side="left", padx=(16, 4))
        self.cb_analytics_focus_cat = ttk.Combobox(triage_bar, values=["All Categories"], state="readonly", width=28)
        self.cb_analytics_focus_cat.pack(side="left", padx=4)
        self.cb_analytics_focus_cat.set("All Categories")
        self.cb_analytics_focus_cat.bind("<<ComboboxSelected>>", lambda e: self.generate_analytics_matrix())

        ttk.Label(triage_bar, text="Min hits:").pack(side="left", padx=(16, 4))
        self.sp_analytics_min_hits = ttk.Spinbox(triage_bar, from_=0, to=999, width=5)
        self.sp_analytics_min_hits.pack(side="left", padx=4)
        self.sp_analytics_min_hits.set(0)
        ttk.Button(triage_bar, text="Apply", command=self.generate_analytics_matrix).pack(side="left", padx=6)
        ttk.Button(triage_bar, text="📤 Export Matrix to CSV", command=self.export_analytics_matrix_csv).pack(
            side="left", padx=(16, 6)
        )

        legend_bar = ttk.Frame(f)
        legend_bar.pack(fill="x", side="top", pady=(0, 6))
        ttk.Label(legend_bar, text="Heat scale:", font=("Segoe UI", 8, "italic")).pack(side="left", padx=(4, 4))
        for step in (0.0, 0.25, 0.5, 0.75, 1.0):
            bg, _fg = _heat_color(step)
            tk.Frame(legend_bar, bg=bg, width=28, height=14, highlightbackground="#cbd5e1", highlightthickness=1).pack(
                side="left", padx=0
            )
        self.lbl_heat_legend_max = ttk.Label(legend_bar, text="  Low  →  High (scale updates with results)", font=("Segoe UI", 8))
        self.lbl_heat_legend_max.pack(side="left", padx=6)

        ttk.Label(
            legend_bar, text="💡 Double-click a title to open its PDF with auto-highlight applied",
            font=("Segoe UI", 8, "italic")
        ).pack(side="right", padx=6)

        container = ttk.Frame(f)
        container.pack(fill="both", expand=True)

        self.analytics_canvas = tk.Canvas(container, bg="#ffffff", highlightthickness=0)
        sb_y = ttk.Scrollbar(container, orient="vertical", command=self.analytics_canvas.yview)
        sb_x = ttk.Scrollbar(container, orient="horizontal", command=self.analytics_canvas.xview)
        self.analytics_canvas.configure(yscrollcommand=sb_y.set, xscrollcommand=sb_x.set)

        sb_x.pack(side="bottom", fill="x")
        sb_y.pack(side="right", fill="y")
        self.analytics_canvas.pack(side="left", fill="both", expand=True)

        self.matrix_grid_frame = tk.Frame(self.analytics_canvas, bg="#ffffff")
        self.canvas_window_id = self.analytics_canvas.create_window((0, 0), window=self.matrix_grid_frame, anchor="nw")

        self.matrix_grid_frame.bind("<Configure>", lambda e: self.analytics_canvas.configure(scrollregion=self.analytics_canvas.bbox("all")))
        self.analytics_canvas.bind("<Configure>", lambda e: self.analytics_canvas.itemconfig(self.canvas_window_id, minwidth=e.width))

    def _category_hit_count(self, haystack: str, cat: str) -> int:
        terms = self.state.keyword_rules.get(cat, {}).get("terms", {})
        active_terms = [t for t, enabled in terms.items() if enabled]
        count = 0
        for t in active_terms:
            t_clean = t.strip().lower()
            if "*" in t_clean:
                prefix = t_clean.replace("*", "")
                count += len(re.findall(r'\b' + re.escape(prefix) + r'\w*', haystack))
            else:
                count += haystack.count(t_clean)
        return count

    def generate_analytics_matrix(self) -> None:
        """Renders headers and rows inside a unified grid frame for 100% stable alignment.

        Rows can be sorted by total relevance (keyword-hit density) and filtered by a
        minimum hit threshold and/or a single focus category, so the user can quickly
        triage which documents are worth opening for deeper analysis.
        """
        for w in self.matrix_grid_frame.winfo_children():
            w.destroy()

        categories = list(self.state.keyword_rules.keys())
        columns = ["Document Title"] + categories

        if not self.active_docs:
            self.lbl_analytics_summary.config(text="Load sources via Control Panel to calculate density matrix.")
            return

        focus_values = ["All Categories"] + categories
        if list(self.cb_analytics_focus_cat["values"]) != focus_values:
            current = self.cb_analytics_focus_cat.get()
            self.cb_analytics_focus_cat.config(values=focus_values)
            self.cb_analytics_focus_cat.set(current if current in focus_values else "All Categories")

        focus_cat = self.cb_analytics_focus_cat.get()
        try:
            min_hits = int(self.sp_analytics_min_hits.get())
        except (ValueError, tk.TclError):
            min_hits = 0

        # Pre-score every document so we can sort/filter before rendering rows.
        scored: list[tuple[Document, dict[str, int], int]] = []
        for doc in self.active_docs:
            haystack = f"{doc.title} {doc.description}".lower()
            per_cat = {cat: self._category_hit_count(haystack, cat) for cat in categories}
            relevance = per_cat.get(focus_cat, sum(per_cat.values())) if focus_cat != "All Categories" else sum(per_cat.values())
            scored.append((doc, per_cat, relevance))

        if min_hits > 0:
            scored = [row for row in scored if row[2] >= min_hits]

        if self.cb_analytics_sort.get() == "Relevance (Total Hits)":
            scored.sort(key=lambda row: row[2], reverse=True)

        if not scored:
            self.lbl_analytics_summary.config(
                text="No documents meet the current filter. Lower 'Min hits' or change the focus category."
            )
            return

        # Configure Grid Weights (using minsize to prevent crashes)
        self.matrix_grid_frame.columnconfigure(0, weight=4, minsize=350)
        for col_idx in range(1, len(columns)):
            self.matrix_grid_frame.columnconfigure(col_idx, weight=1, minsize=110)

        # Render Header Row
        for col_idx, heading_text in enumerate(columns):
            is_title = (col_idx == 0)
            lbl = tk.Label(
                self.matrix_grid_frame,
                text=heading_text,
                font=("Segoe UI", 9, "bold"),
                bg="#00247d" if is_title else "#1e293b",
                fg="#ffffff",
                anchor="w" if is_title else "center",
                padx=10,
                pady=10,
                justify="left" if is_title else "center",
                wraplength=320 if is_title else 100,
                relief="ridge",
                bd=1,
            )
            lbl.grid(row=0, column=col_idx, sticky="nsew")

        cat_totals = {cat: 0 for cat in categories}
        total_occurrences = 0

        # Global max across every visible cell, used to scale the heat gradient
        # so color intensity is comparable across the whole matrix at a glance.
        global_max = max((max(per_cat.values()) for _doc, per_cat, _rel in scored if per_cat), default=0)
        self.lbl_heat_legend_max.config(
            text=f"  Low (1)  →  High ({global_max}) — darkest red = most keyword-dense cell"
        )

        # Render Data Rows (from the pre-scored/sorted/filtered list)
        for row_idx, (doc, per_cat, _relevance) in enumerate(scored, start=1):
            doc_matches = 0
            row_bg = "#f8fafc" if row_idx % 2 == 0 else "#ffffff"
            real_idx = self.active_docs.index(doc)

            # Title Cell
            lbl_title = tk.Label(
                self.matrix_grid_frame,
                text=doc.title,
                font=("Segoe UI", 9),
                bg=row_bg,
                fg="#0f172a",
                anchor="w",
                padx=10,
                pady=8,
                justify="left",
                wraplength=350,
                relief="solid",
                bd=1,
            )
            lbl_title.grid(row=row_idx, column=0, sticky="nsew")
            lbl_title.bind("<Double-1>", lambda e, idx=real_idx: self.jump_to_reader_doc(idx, auto_open_pdf=True))

            # Category Count Cells
            for col_idx, cat in enumerate(categories, start=1):
                count = per_cat[cat]
                cat_totals[cat] += count
                doc_matches += count

                if count > 0 and global_max > 0:
                    cell_bg, cell_fg = _heat_color(count / global_max)
                else:
                    cell_bg, cell_fg = row_bg, "#94a3b8"

                lbl_val = tk.Label(
                    self.matrix_grid_frame,
                    text=str(count) if count > 0 else "—",
                    font=("Segoe UI", 9, "bold" if count > 0 else "normal"),
                    bg=cell_bg,
                    fg=cell_fg,
                    anchor="center",
                    padx=6,
                    pady=8,
                    relief="solid",
                    bd=1,
                )
                lbl_val.grid(row=row_idx, column=col_idx, sticky="nsew")

            total_occurrences += doc_matches

        # Totals footer row - quick sense of which categories dominate this
        # whole result set, independent of any single document.
        footer_row = len(scored) + 1
        lbl_footer_title = tk.Label(
            self.matrix_grid_frame, text="📊 Column Totals", font=("Segoe UI", 9, "bold"),
            bg="#1e293b", fg="#ffffff", anchor="w", padx=10, pady=8, relief="ridge", bd=1,
        )
        lbl_footer_title.grid(row=footer_row, column=0, sticky="nsew")
        for col_idx, cat in enumerate(categories, start=1):
            total_for_cat = cat_totals[cat]
            lbl_footer_val = tk.Label(
                self.matrix_grid_frame, text=str(total_for_cat), font=("Segoe UI", 9, "bold"),
                bg="#1e293b", fg="#ffffff", anchor="center", padx=6, pady=8, relief="ridge", bd=1,
            )
            lbl_footer_val.grid(row=footer_row, column=col_idx, sticky="nsew")

        # Stash for CSV export so the export matches exactly what's on screen.
        self._last_analytics_export = {"categories": categories, "scored": scored, "cat_totals": cat_totals}

        num_docs = len(scored)
        num_total = len(self.active_docs)
        top_cat = max(cat_totals, key=cat_totals.get) if cat_totals else "None"
        top_count = cat_totals.get(top_cat, 0)
        shown = f"{num_docs} of {num_total}" if num_docs != num_total else f"{num_docs}"
        self.lbl_analytics_summary.config(
            text=f"📄 {shown} Documents Shown  |  🎯 Total Keyword Hits: {total_occurrences}  |  🔥 Dominant Domain: '{top_cat}' ({top_count} hits)"
        )

    def export_analytics_matrix_csv(self) -> None:
        """Exports the currently-displayed (sorted/filtered) matrix, including the
        column-totals row, so the analysis can continue in a spreadsheet."""
        export_data = getattr(self, "_last_analytics_export", None)
        if not export_data or not export_data.get("scored"):
            messagebox.showinfo("Nothing to Export", "Load sources and generate the matrix first.")
            return

        out_p = filedialog.asksaveasfilename(defaultextension=".csv", filetypes=[("CSV File", "*.csv")])
        if not out_p:
            return

        import csv as _csv

        categories = export_data["categories"]
        try:
            with open(out_p, "w", newline="", encoding="utf-8") as fh:
                writer = _csv.writer(fh)
                writer.writerow(["Document Title", "URL"] + categories + ["Total Hits"])
                for doc, per_cat, relevance in export_data["scored"]:
                    writer.writerow([doc.title, doc.url] + [per_cat[c] for c in categories] + [sum(per_cat.values())])
                writer.writerow(
                    ["Column Totals", ""] + [export_data["cat_totals"][c] for c in categories] + [""]
                )
        except OSError as exc:
            messagebox.showerror("Export Failed", f"Could not write file:\n{exc}")
            return
        messagebox.showinfo("Export Complete", f"Saved to: {out_p}")

    def jump_to_reader_doc(self, doc_idx: int, auto_open_pdf: bool = False) -> None:
        if not (0 <= doc_idx < len(self.active_docs)):
            return
        self.lb_docs.selection_clear(0, tk.END)
        self.lb_docs.selection_set(doc_idx)
        self.lb_docs.see(doc_idx)
        self.on_doc_select(None)
        self.nb.select(self.tab_reader)

        if not auto_open_pdf:
            return

        doc = self.active_docs[doc_idx]
        pdf_url = next(
            (url for url in doc.attachments if "📕 PDF" in classify_attachment_url(url)), None
        )
        if pdf_url is None:
            return
        self._open_pdf_attachment(
            pdf_url, self.reader_pdf_viewer, self.reader_nb,
            on_loaded=lambda: self.reader_pdf_viewer.auto_highlight(show_dialog=False),
        )

    # ======================================================================
    # Dialogs / misc tools
    # ======================================================================
    def open_split_comparator(self) -> None:
        SplitPDFComparatorWindow(self.root, lambda: self.state.keyword_rules)

    def open_history_dialog(self) -> None:
        pop = tk.Toplevel(self.root)
        pop.title("📜 Search History")
        pop.geometry("600x400")
        lb = tk.Listbox(pop, bg="#ffffff", font=("Segoe UI", 10))
        lb.pack(fill="both", expand=True, padx=10, pady=10)
        for rec in self.state.search_history:
            lb.insert(tk.END, f"[{rec.get('timestamp')}] Query: '{rec.get('query')}' | Dept: {rec.get('dept')}")

        def rerun(_event):
            sel = lb.curselection()
            if sel:
                self.e_topic.delete(0, tk.END)
                self.e_topic.insert(0, self.state.search_history[sel[0]].get("query"))
                pop.destroy()
                self.start_harvest()

        lb.bind("<Double-1>", rerun)

    def copy_citation_popup(self) -> None:
        if not self.selected_doc:
            return
        year = datetime.now().strftime("%Y")
        apa = f"HM Government. ({year}). {self.selected_doc.title}. GOV.UK."
        oscola = f"HM Government, '{self.selected_doc.title}' ({year}) accessed {datetime.now().strftime('%B %Y')}."
        pop = tk.Toplevel(self.root)
        pop.title("📋 Reference Generator")
        pop.geometry("550x250")
        ttk.Label(pop, text="APA Format:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=(10, 2))
        t1 = tk.Text(pop, height=2, wrap="word")
        t1.pack(fill="x", padx=10)
        t1.insert(tk.END, apa)
        ttk.Label(pop, text="OSCOLA Format:", font=("Segoe UI", 9, "bold")).pack(anchor="w", padx=10, pady=(10, 2))
        t2 = tk.Text(pop, height=2, wrap="word")
        t2.pack(fill="x", padx=10)
        t2.insert(tk.END, oscola)

    def export_word_briefing(self) -> None:
        if not HAS_DOCX:
            messagebox.showwarning("Library Missing", "Run 'pip install python-docx' first.")
            return
        briefing_text = self.txt_briefing.get("1.0", tk.END).strip()
        if not briefing_text:
            messagebox.showinfo("No Briefing", "Please run a search scan first.")
            return
        doc = docx.Document()
        doc.add_heading(f"GOV.UK Briefing: {self.e_topic.get().strip().upper()}", 0)
        doc.add_heading("Executive Summary", level=1)
        doc.add_paragraph(briefing_text)
        out_p = filedialog.asksaveasfilename(defaultextension=".docx", filetypes=[("Word Document", "*.docx")])
        if out_p:
            doc.save(out_p)
            messagebox.showinfo("Export Complete", f"Saved to: {out_p}")

    def cross_pdf_search_dialog(self) -> None:
        pop = tk.Toplevel(self.root)
        pop.title("🔍 Multi-PDF Cross Search")
        pop.geometry("600x450")
        e = ttk.Entry(pop)
        e.pack(fill="x", padx=10, pady=5)
        txt = scrolledtext.ScrolledText(pop, wrap="word")
        txt.pack(fill="both", expand=True, padx=10, pady=5)
        status = ttk.Label(pop, text="")
        status.pack(fill="x", padx=10)

        def run():
            term = e.get().strip().lower()
            if not term:
                return
            txt.delete("1.0", tk.END)
            status.config(text="Searching…")
            threading.Thread(target=self._cross_pdf_search_worker, args=(term, txt, status), daemon=True).start()

        ttk.Button(pop, text="Search PDFs", command=run).pack(pady=5)

    def _cross_pdf_search_worker(self, term: str, txt: scrolledtext.ScrolledText, status: ttk.Label) -> None:
        import pypdfium2 as pdfium

        matches_found = 0
        for r, _dirs, files in os.walk(config.DATA_DIR):
            for file in files:
                if not file.endswith(".pdf"):
                    continue
                try:
                    doc = pdfium.PdfDocument(os.path.join(r, file))
                except Exception as exc:
                    logger.info("Skipping unreadable PDF %s: %s", file, exc)
                    continue
                for p_num, page in enumerate(doc):
                    textpage = page.get_textpage()
                    if term in textpage.get_text_range().lower():
                        matches_found += 1
                        line = f"📌 [File: {file} | Page {p_num + 1}]\n"
                        self.root.after(0, lambda line=line: txt.insert(tk.END, line))
        self.root.after(0, lambda: status.config(text=f"Done — {matches_found} match(es) found."))